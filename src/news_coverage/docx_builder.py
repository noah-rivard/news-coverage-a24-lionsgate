"""DOCX generation for multi-buyer news coverage reports."""

from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Dict, List, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


SECTION_TITLES: List[Tuple[str, str]] = [
    ("Highlights", "Highlights From The Quarter"),
    ("Org", "Org"),
    ("Content / Deals / Distribution", "Content, Deals & Distribution"),
    ("Strategy & Miscellaneous News", "Strategy & Miscellaneous News"),
    ("Investor Relations", "Investor Relations"),
    ("M&A", "M&A"),
]

MEDIUM_ORDER = ("Film", "TV", "Specials", "International", "Sports/Podcasts", "General")


@dataclass
class CoverageEntry:
    title: str
    url: str
    published_at: date
    section: str
    subheading: str | None
    medium: str
    summary_lines: List[str]


@dataclass
class BuyerReport:
    buyer: str
    entries: List[CoverageEntry] = field(default_factory=list)


def _month_range_text(quarter_label: str) -> str:
    mapping = {
        "Q1": "January – March",
        "Q2": "April – June",
        "Q3": "July – September",
        "Q4": "October – December",
    }
    try:
        q = quarter_label.strip().split()[1]
    except Exception:
        return "January – March"
    return mapping.get(q, "January – March")


def _format_md(dt: date) -> str:
    return f"{dt.month}/{dt.day}"


def _group_entries(
    entries: Sequence[CoverageEntry],
) -> Dict[str, Dict[str, Dict[str, List[CoverageEntry]]]]:
    """
    Group entries by section -> medium -> subheading, sorted newest to oldest.
    Returns nested dict: {section: {medium: {subheading: [entries...]}}}
    """
    grouped: Dict[str, Dict[str, Dict[str, List[CoverageEntry]]]] = defaultdict(
        lambda: defaultdict(lambda: defaultdict(list))
    )
    for e in entries:
        sub = (e.subheading or "General News & Strategy").strip() or "General News & Strategy"
        grouped[e.section][e.medium][sub].append(e)

    for section_map in grouped.values():
        for medium_map in section_map.values():
            for subheading, items in medium_map.items():
                medium_map[subheading] = sorted(
                    items, key=lambda i: i.published_at, reverse=True
                )
    return grouped


def _set_title_styles(doc: Document) -> None:
    """
    Apply minimal styling to approximate the reference docs without embedding custom themes.
    Keeps ASCII; uses Calibri 14/12 for header/subheader for readability.
    """
    title_style = doc.styles["Title"]
    title_font = title_style.font
    title_font.name = "Calibri"
    title_font.size = Pt(20)

    try:
        title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass


def build_docx(report: BuyerReport, output_path: Path, quarter_label: str) -> None:
    """Render a single buyer DOCX."""
    doc = Document()
    _set_title_styles(doc)

    def _safe_add_paragraph(*, text: str = "", style: str | None = None):
        if style is None:
            return doc.add_paragraph(text)
        try:
            return doc.add_paragraph(text, style=style)
        except KeyError:
            return doc.add_paragraph(text)

    def _safe_add_heading(text: str, level: int):
        try:
            return doc.add_heading(text, level=level)
        except KeyError:
            return doc.add_heading(text, level=level)

    def _ordered_subheadings(sub_map: Dict[str, List[CoverageEntry]]) -> List[str]:
        preferred = [
            "General News & Strategy",
            "Development",
            "Pickups",
            "Dating",
            "Greenlights",
            "Renewals",
            "Cancellations",
        ]
        preferred_set = set(preferred)
        ordered = [s for s in preferred if s in sub_map]
        extras = sorted([s for s in sub_map.keys() if s not in preferred_set])
        return ordered + extras

    # Cover/header
    title = doc.add_heading(f"{quarter_label} News & Updates", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"{_month_range_text(quarter_label)} {quarter_label.split()[0]}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    grouped = _group_entries(report.entries)

    for idx, (section_key, section_display) in enumerate(SECTION_TITLES, start=0):
        if section_key not in grouped:
            continue
        doc.add_heading(f"{idx}. {section_display}", level=1)
        section_entries = grouped[section_key]

        # Only Content & Strategy get medium grouping emphasis
        if section_key in {"Content / Deals / Distribution", "Strategy & Miscellaneous News"}:
            medium_order = MEDIUM_ORDER
        else:
            medium_order = ("General",)

        for medium in medium_order:
            if medium not in section_entries:
                continue
            if section_key in {"Content / Deals / Distribution", "Strategy & Miscellaneous News"}:
                _safe_add_heading(medium, level=2)
                medium_map = section_entries[medium]
                for subheading in _ordered_subheadings(medium_map):
                    sub_para = _safe_add_paragraph(text=subheading, style="No Spacing")
                    if sub_para.runs:
                        sub_para.runs[0].bold = True
                    else:
                        sub_para.add_run(subheading).bold = True

                    for entry in medium_map[subheading]:
                        line_para = _safe_add_paragraph(style="No Spacing")
                        if ":" in entry.title:
                            left, right = entry.title.split(":", 1)
                            label = left.strip()
                            remainder = right.lstrip()
                            if label.lower() in {"interview", "commentary"}:
                                label_run = line_para.add_run(f"{label}:")
                                label_run.italic = True
                                remainder_run = line_para.add_run(
                                    f" {remainder} ({_format_md(entry.published_at)})"
                                )
                                remainder_run.bold = True
                            else:
                                label_run = line_para.add_run(f"{label}:")
                                label_run.bold = True
                                label_run.italic = True
                                line_para.add_run(
                                    f" {remainder} ({_format_md(entry.published_at)})"
                                )
                        else:
                            line_para.add_run(
                                f"{entry.title} ({_format_md(entry.published_at)})"
                            )

                        for line in entry.summary_lines:
                            if (line or "").strip():
                                _safe_add_paragraph(text=line.strip(), style="No Spacing")
            else:
                medium_map = section_entries[medium]
                for subheading, entries in medium_map.items():
                    if section_key != "Highlights":
                        _safe_add_heading(subheading, level=2)
                    for entry in entries:
                        bullet = _safe_add_paragraph(style="List Paragraph")
                        date_text = _format_md(entry.published_at)
                        inline_note = ""
                        remaining_lines = list(entry.summary_lines)
                        if (
                            section_key == "Org"
                            and (subheading or "").strip() == "Exec Changes"
                            and remaining_lines
                        ):
                            inline_note = remaining_lines[0].strip()
                            remaining_lines = remaining_lines[1:]

                        if inline_note:
                            bullet.add_run(f"{entry.title} ({date_text}) {inline_note}")
                        else:
                            bullet.add_run(f"{entry.title} ({date_text})")

                        for line in remaining_lines:
                            if not (line or "").strip():
                                continue
                            _safe_add_paragraph(
                                text=line.strip(),
                                style="List Paragraph",
                            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
