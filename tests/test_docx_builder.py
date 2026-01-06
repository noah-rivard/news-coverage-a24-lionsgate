from datetime import date

from docx import Document

from news_coverage.docx_builder import BuyerReport, CoverageEntry, build_docx


def test_build_docx_renders_content_list_entries_with_optional_note(tmp_path):
    report = BuyerReport(
        buyer="WBD",
        entries=[
            CoverageEntry(
                title="Wuthering Heights: Warner Bros. Pictures, drama",
                url="https://example.com/wuthering",
                published_at=date(2024, 10, 24),
                section="Content / Deals / Distribution",
                subheading="Pickups",
                medium="Film",
                summary_lines=[
                    "The studio won the package by committing a healthy P&A spend."
                ],
            ),
            CoverageEntry(
                title="Animal Friends: Warner Bros. Pictures, road trip adventure",
                url="https://example.com/animal",
                published_at=date(2024, 12, 9),
                section="Content / Deals / Distribution",
                subheading="Pickups",
                medium="Film",
                summary_lines=[],
            ),
        ],
    )

    output_path = tmp_path / "out.docx"
    build_docx(report, output_path, quarter_label="2024 Q4")

    doc = Document(output_path)
    texts = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]

    assert "Pickups" in texts
    assert any(t.startswith("Wuthering Heights:") for t in texts)
    assert "The studio won the package by committing a healthy P&A spend." in texts
    assert any(t.startswith("Animal Friends:") for t in texts)

    wuthering_para = next(
        p for p in doc.paragraphs if (p.text or "").strip().startswith("Wuthering Heights:")
    )
    assert wuthering_para.runs
    assert wuthering_para.runs[0].bold is True
    assert wuthering_para.runs[0].italic is True


def test_build_docx_renders_interview_header_and_paragraphs(tmp_path):
    report = BuyerReport(
        buyer="WBD",
        entries=[
            CoverageEntry(
                title=(
                    "Interview: Kathleen Finch, exiting Chairman and CEO of US Networks "
                    "at Warner Bros. Discovery"
                ),
                url="https://example.com/interview",
                published_at=date(2024, 10, 24),
                section="Content / Deals / Distribution",
                subheading="General News & Strategy",
                medium="TV",
                summary_lines=[
                    "Finch played a key role in shaping content strategy at Food Network and HGTV.",
                    (
                        "Her successor, Channing Dungey, will inherit a strategy focused on "
                        "maximizing audience reach."
                    ),
                ],
            )
        ],
    )

    output_path = tmp_path / "out.docx"
    build_docx(report, output_path, quarter_label="2024 Q4")

    doc = Document(output_path)
    texts = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]

    assert "General News & Strategy" in texts
    assert any(t.startswith("Interview: Kathleen Finch") for t in texts)
    assert (
        "Finch played a key role in shaping content strategy at Food Network and HGTV."
        in texts
    )
    assert (
        "Her successor, Channing Dungey, will inherit a strategy focused on "
        "maximizing audience reach."
        in texts
    )

    interview_para = next(
        p for p in doc.paragraphs if (p.text or "").strip().startswith("Interview:")
    )
    assert len(interview_para.runs) >= 2
    assert interview_para.runs[0].italic is True
    assert interview_para.runs[1].bold is True


def test_build_docx_renders_exec_change_note_inline_after_date(tmp_path):
    note = (
        "Brett will continue to report to Channing Dungey, Warner Bros. Television Group "
        "chairman/CEO."
    )
    report = BuyerReport(
        buyer="WBD",
        entries=[
            CoverageEntry(
                title="Promotion: Brett Paul, COO of U.S. Networks at Warner Bros. Discovery",
                url="https://example.com/promo",
                published_at=date(2024, 12, 18),
                section="Org",
                subheading="Exec Changes",
                medium="General",
                summary_lines=[
                    note,
                ],
            )
        ],
    )

    output_path = tmp_path / "out.docx"
    build_docx(report, output_path, quarter_label="2024 Q4")

    doc = Document(output_path)
    texts = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]

    assert any(
        t.startswith("Promotion: Brett Paul")
        and "(12/18)" in t
        and "Brett will continue to report" in t
        for t in texts
    )
    assert note not in texts
